import streamlit as st
import g4f
from g4f.client import Client
from docx import Document
import os
import pandas as pd
import docx
import re
import time
import random
from io import BytesIO

# Definir la función de chat
def chatear(modelo, mensaje):
    client = Client()
    with st.spinner('Contactando a la AI ...'):
        response = client.chat.completions.create(
            model=modelo,
            
            messages=[{'role': 'user', 'content': mensaje}]
        )
        respuesta = response.choices[0].message.content
    return respuesta

# Definir la función para extraer recomendaciones
def extraer_recomendaciones(texto):
    texto = texto.strip("'")
    patron = r'\d+\) [^0-9]+(?=(?:\d+\)|$))'
    coincidencias = re.findall(patron, texto)
    recomendaciones = [coincidencia.strip() for coincidencia in coincidencias]
    return recomendaciones

def main():
    # Configurar la interfaz de Streamlit
    st.title('Generador de Cursos asistido por I.A')
    # Entrada de texto para el nombre del curso
    idioma = st.selectbox("Selecciona el idioma", ["Español", "English","Portuguese","Italian","French"])
    modelo = st.selectbox("Selecciona el modelo de lenguaje", ["gpt-4-0613", "gpt-3.5-turbo","gpt-4o","gpt-4-turbo"])
    nombre = st.text_input("Escribe el nombre del curso:")
    # Entrada de texto para el público objetivo
    publico = st.text_input("¿A quién está dirigido este curso?")
    # Entrada numérica para la cantidad de módulos
    cant_modulos = st.number_input("¿En cuántos módulos quiere dividir el curso?", min_value=1, step=1)

    # Botón para generar la descripción del curso
    if st.button('Generar Descripción'):
        if nombre and publico and cant_modulos:

            texto = f'Nombre del curso: {nombre} - Público objetivo: {publico}'
            metas = f'Actúa como un especialista en pedagogía y educación a distancia (elearning). Debes crear la descripción de un curso en base a: {texto} . Al final debes incluir los objetivos de aprendizaje. Recuerda hacerlo en idioma {idioma}.'

            st.write('Espere ... contactando a ChatGPT')
            respuesta = chatear(modelo, metas)
            st.write('Descripción generada:')
            descripcion=respuesta
            st.write(respuesta)

            recomendaciones = extraer_recomendaciones(respuesta)
            st.write('Recomendaciones extraídas:')
            for recomendacion in recomendaciones:
                st.write(recomendacion)

            st.write(f'Guardando descripción de {nombre}')
            
            arma_modulos = f'Actúa como un especialista en pedagogía y educación a distancia (elearning). Debes crear y enumerar {cant_modulos} módulos para impartir el curso: {texto}. Debes enumerar los mismos numéricamente así: 1)texto, 2)texto, etc. Solo quiero que me des los módulos sin generar descripciones y sin saltos de línea entre ellos.Recuerda hacerlo en idioma {idioma}.'
            st.write('Espere. Contactando a ChatGPT para crear módulos ... ')
            respuesta = chatear(modelo, arma_modulos)
            st.write('Módulos finalizados ... ')
            elementos = respuesta

            # Extraer recomendaciones
            recomendaciones = extraer_recomendaciones(elementos)
            st.write(recomendaciones)
            df_modulos = pd.DataFrame(recomendaciones, columns=["modulos"])
            df_modulos['contenido'] = 'NaN'

            # Mostrar DataFrame inicial
            st.write('Módulos generados:')
            st.write(df_modulos)
            
            # Agregar contenido para cada módulo
            for i in df_modulos.index:
                arma_modulos = f'Actúa como un especialista en pedagogía y educación a distancia (elearning). Debes crear el contenido para el módulo {df_modulos.modulos[i]} del curso: {nombre}. Al final debes incluir fuentes con información complementaria bajo el título: Puedes encontrar más información en los siguientes links. Recuerda hacerlo en idioma {idioma}. '
                st.write(f'Generando contenido para el módulo: {df_modulos.modulos[i]}')
                respuesta = chatear(modelo, arma_modulos)
                df_modulos.at[i, 'contenido'] = respuesta
                time.sleep(10)  # Esperar 30 segundos entre solicitudes para evitar sobrecargar el servidor
            
            # Mostrar DataFrame final con contenido
            st.write(df_modulos)
            
             # Agregar cuestionario para cada módulo
            df_modulos['cuestionario'] = 'NaN'

            for i in df_modulos.index:
                choice = f'Actúa como un especialista en pedagogía y educación a distancia (elearning). Debes crear un cuestionario multiple choice de 10 preguntas (indicando con una X la opción correcta) teniendo en cuenta el contenido y específicamente para el módulo {df_modulos.modulos[i]} del curso: {nombre}. El ejercicio es individual.Recuerda hacerlo en idioma {idioma}.'
                simulacion = f'Actúa como un experto en elearning y pedagogía. Eres el docente de un curso virtual de {nombre} del módulo {df_modulos.modulos[i]} y tienes que realizar una actividad de simulación individual donde el alumno debe tomar un rol, evaluar una situación y poner en práctica lo aprendido. Debe incluir ejercicios que evalúen la toma de decisiones en situaciones particulares o preguntar con qué otros roles debería interactuar para poder cumplir la asignación. El ejercicio es individual.Recuerda hacerlo en idioma {idioma}.'
                completar = f'Actúa como un experto en elearning y pedagogía. Eres el docente de un curso virtual de {nombre} del módulo {df_modulos.modulos[i]} y tienes que realizar una actividad de simulación individual donde el alumno debe tomar un rol, evaluar una situación y poner en práctica lo aprendido. Debe incluir ejercicios como completar una oración. El ejercicio es individual.Recuerda hacerlo en idioma {idioma}.'
                ejercicios = [choice, simulacion, completar]
                ejercicio_random = random.choice(ejercicios)
                
                st.write(f'Generando Ejercicios aleatorios para el módulo: {i + 1}')
                respuesta = chatear(modelo, ejercicio_random)
                df_modulos.at[i, 'cuestionario'] = respuesta
                time.sleep(10)

            # Mostrar DataFrame final con cuestionarios
            st.write(df_modulos)
            
            # Crear el documento Word
            elementos = str(elementos).replace(". ", ".\n ")
            document = Document()
            document.add_heading(nombre, 0)
            document.add_heading('Descripción:', level=1)
            document.add_paragraph(descripcion)
            document.add_heading('Módulos:', level=1)
            document.add_paragraph(elementos)

            for i in df_modulos.index:
                document.add_heading('Módulo ' + str(i + 1), level=1)
                document.add_paragraph(df_modulos['modulos'][i])
                document.add_heading('Contenido:', level=1)
                document.add_paragraph(df_modulos['contenido'][i])
                document.add_heading('Cuestionario:', level=1)
                document.add_paragraph(df_modulos['cuestionario'][i])
            
            # Guardar el documento en un BytesIO para descarga
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Descargar curso en Word",
                data=buffer,
                file_name="curso.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            
        else:
            st.error('Por favor, complete todos los campos.')

if __name__ == "__main__":
    main()